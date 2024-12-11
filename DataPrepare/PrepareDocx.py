from docx import Document

class DocxParser:
    @classmethod
    def parse_docx(cls, file_path: str) -> list:
        doc = Document(file_path)
        result = []
        tables_iter = iter(doc.tables)
        current_list = []

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name.lower()
            text = paragraph.text.strip()

            if "list paragraph" in style_name:
                if text:
                    current_list.append(text)
                continue
            else:
                if current_list:
                    list_text = " ".join(current_list)
                    result.append({"type": "list", "text": list_text})
                    current_list = []

            if not text:
                continue

            result.append({"type": "paragraph", "text": text})

            try:
                if len(paragraph._element.xpath("./following-sibling::w:tbl")) > 0:
                    table = next(tables_iter)
                    table_data = {}

                    # Инициализация столбцов
                    for row in table.rows:
                        for i, cell in enumerate(row.cells):
                            column_name = f"Column {i+1}"

                            # Инициализируем список, если этого еще не сделали
                            if column_name not in table_data:
                                table_data[column_name] = []

                            # Добавляем данные в соответствующий столбец
                            cell_text = cell.text.strip()
                            if cell_text:
                                table_data[column_name].append(cell_text)

                    # Преобразуем данные в нужный формат
                    table_text = "; ".join([f"{column}: {', '.join(values)}" for column, values in table_data.items()])
                    result.append({"type": "table", "text": table_text})
            except StopIteration:
                pass

        if current_list:
            list_text = " ".join(current_list)
            result.append({"type": "list", "text": list_text})

        return result
