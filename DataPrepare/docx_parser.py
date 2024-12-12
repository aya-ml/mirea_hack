from docx import Document

class DocxParser:
    def __init__(self, chunk_size=50):
        self.chunk_size = chunk_size
        self.global_chunk_id = 0

    def parse_docx(self, file_path: str, document_id: str) -> list:
        doc = Document(file_path)
        result = []
        tables_iter = iter(doc.tables)
        current_list = []
        current_chunk = []

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name.lower()
            text = paragraph.text.strip()

            if "list paragraph" in style_name:
                if text:
                    current_list.append(text)
                continue
            else:
                if current_list:
                    list_text = " ".join(current_list)
                    self._add_chunks(result, list_text, document_id)
                    current_list = []

            if not text:
                continue

            current_chunk.extend(text.split())
            if len(current_chunk) >= self.chunk_size:
                chunk_text = " ".join(current_chunk[:self.chunk_size])
                self._add_chunk(result, chunk_text, document_id)
                current_chunk = current_chunk[self.chunk_size:]

            self._add_chunk(result, text, document_id)

            try:
                if len(paragraph._element.xpath("./following-sibling::w:tbl")) > 0:
                    table = next(tables_iter)
                    table_text = self._format_table(table)
                    self._add_chunk(result, table_text, document_id)
            except StopIteration:
                pass

        if current_list:
            list_text = " ".join(current_list)
            self._add_chunks(result, list_text, document_id)

        if current_chunk:
            chunk_text = " ".join(current_chunk)
            self._add_chunk(result, chunk_text, document_id)

        return result

    def _format_table(self, table):
        table_lines = []

        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_lines.append(" | ".join(row_data))

        return "\n".join(table_lines)

    def _add_chunk(self, result, text, document_id):
        chunk_id = f"{document_id}_{self.global_chunk_id}"
        result.append({"text": text, "chunk_id": chunk_id})
        self.global_chunk_id += 1

    def _add_chunks(self, result, text, document_id):
        words = text.split()
        for i in range(0, len(words), self.chunk_size):
            chunk_text = " ".join(words[i:i + self.chunk_size])
            self._add_chunk(result, chunk_text, document_id)
